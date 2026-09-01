import os
import json
import re
import subprocess
import tempfile
import base64
from pathlib import Path

import streamlit as st
import streamlit.components.v1 as components
import whisper
import torch
import soundfile as sf
from pyannote.audio import Pipeline
from deep_translator import GoogleTranslator
from docx import Document

from reportlab.lib.pagesizes import A4
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER


# =========================================================
# PAGE CONFIGURATION
# =========================================================

st.set_page_config(
    page_title="AI Video Transcription Tool",
    page_icon="🎙️",
    layout="wide"
)


# =========================================================
# SESSION STATE
# =========================================================

DEFAULT_STATE = {
    "generated": False,
    "transcript": "",
    "segments": [],
    "detected_language": "",
    "model_name": "",
    "duration": None,
    "video_name": "",
    "edited_transcript": ""
}

for key, value in DEFAULT_STATE.items():

    if key not in st.session_state:
        st.session_state[key] = value


# =========================================================
# HEADER
# =========================================================

st.title("🎙️ AI Video Transcription Tool")

st.write(
    "Upload a video and automatically convert spoken content "
    "into a timestamped transcript."
)

st.info(
    "🔒 Your video is processed locally on this computer. "
    "No paid API is required."
)


# =========================================================
# DEVICE
# =========================================================

def get_device():

    try:

        import torch

        if torch.cuda.is_available():
            return "cuda"

        return "cpu"

    except Exception:

        return "cpu"


device = get_device()

@st.cache_resource
def load_diarization_pipeline():
    pipeline = Pipeline.from_pretrained(
        "pyannote/speaker-diarization-community-1"
    )

    if torch.cuda.is_available():
        pipeline.to(torch.device("cuda"))

    return pipeline


def perform_speaker_diarization(audio_path, num_speakers=None):
    audio, sample_rate = sf.read(audio_path)

    waveform = torch.from_numpy(audio).float()

    if waveform.ndim == 1:
        waveform = waveform.unsqueeze(0)
    else:
        waveform = waveform.transpose(0, 1)

    pipeline = load_diarization_pipeline()
    if num_speakers is not None:
        diarization = pipeline(
            {"waveform": waveform, "sample_rate": sample_rate},
            num_speakers=num_speakers
        )
    else:
        diarization = pipeline(
            {"waveform": waveform, "sample_rate": sample_rate}
        )

    return diarization

def add_speakers_to_segments(segments, diarization):
    updated_segments = []

    speaker_map = {}
    speaker_count = 0

    for segment in segments:
        start = float(segment.get("start", 0))
        end = float(segment.get("end", 0))

        speaker_overlap = {}

        for turn, _, speaker in diarization.speaker_diarization.itertracks(
            yield_label=True
        ):
            overlap_start = max(start, turn.start)
            overlap_end = min(end, turn.end)
            overlap = max(0, overlap_end - overlap_start)

            if overlap > 0:
                speaker_overlap[speaker] = (
                    speaker_overlap.get(speaker, 0) + overlap
                )

        if speaker_overlap:
            detected_speaker = max(
                speaker_overlap,
                key=speaker_overlap.get
            )

            if detected_speaker not in speaker_map:
                speaker_count += 1
                speaker_map[detected_speaker] = (
                    f"Speaker {speaker_count}"
                )

            speaker_label = speaker_map[detected_speaker]

        else:
            speaker_label = "UNKNOWN"

        updated_segment = segment.copy()
        updated_segment["speaker"] = speaker_label

        updated_segments.append(updated_segment)

    return updated_segments

if device == "cuda":

    st.success(
        "⚡ CUDA GPU detected — GPU acceleration enabled."
    )

else:

    st.warning(
        "💻 CUDA GPU not detected — processing will use CPU."
    )


# =========================================================
# WHISPER MODELS
# =========================================================

MODEL_OPTIONS = {
    "Tiny — Fastest": "tiny",
    "Base — Balanced": "base",
    "Small — Better Accuracy": "small",
    "Medium — High Accuracy": "medium"
}


@st.cache_resource
def load_whisper_model(
    model_name,
    device_name
):

    with st.spinner(
        f"Loading Whisper '{model_name}' model..."
    ):

        return whisper.load_model(
            model_name,
            device=device_name
        )


# =========================================================
# LANGUAGE OPTIONS
# =========================================================

LANGUAGE_OPTIONS = {

    "Auto Detect": None,

    "English": "en",
    "Hindi": "hi",
    "Urdu": "ur",
    "Bengali": "bn",
    "Marathi": "mr",
    "Gujarati": "gu",

    "Tamil": "ta",
    "Telugu": "te",
    "Kannada": "kn",
    "Malayalam": "ml",
    "Punjabi": "pa",

    "Arabic": "ar",
    "Spanish": "es",
    "French": "fr",
    "German": "de",
    "Portuguese": "pt",
    "Russian": "ru",

    "Japanese": "ja",
    "Korean": "ko",
    "Chinese": "zh"
}


# =========================================================
# TIME FUNCTIONS
# =========================================================

def format_display_time(seconds):

    seconds = float(seconds)

    hours = int(seconds // 3600)

    minutes = int(
        (seconds % 3600) // 60
    )

    secs = int(seconds % 60)

    if hours > 0:

        return (
            f"{hours:02d}:"
            f"{minutes:02d}:"
            f"{secs:02d}"
        )

    return (
        f"{minutes:02d}:"
        f"{secs:02d}"
    )


def format_srt_time(seconds):

    seconds = float(seconds)

    hours = int(seconds // 3600)

    minutes = int(
        (seconds % 3600) // 60
    )

    secs = int(seconds % 60)

    milliseconds = int(
        (seconds - int(seconds)) * 1000
    )

    return (
        f"{hours:02d}:"
        f"{minutes:02d}:"
        f"{secs:02d},"
        f"{milliseconds:03d}"
    )


def format_vtt_time(seconds):

    seconds = float(seconds)

    hours = int(seconds // 3600)

    minutes = int(
        (seconds % 3600) // 60
    )

    secs = int(seconds % 60)

    milliseconds = int(
        (seconds - int(seconds)) * 1000
    )

    return (
        f"{hours:02d}:"
        f"{minutes:02d}:"
        f"{secs:02d}."
        f"{milliseconds:03d}"
    )


# =========================================================
# VIDEO DURATION
# =========================================================

def get_video_duration(video_path):

    try:

        command = [
            "ffprobe",
            "-v",
            "error",
            "-show_entries",
            "format=duration",
            "-of",
            "default=noprint_wrappers=1:nokey=1",
            video_path
        ]

        result = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=30
        )

        if result.returncode == 0:

            return float(
                result.stdout.strip()
            )

    except Exception:

        pass

    return None


# =========================================================
# TRANSCRIPT
# =========================================================

def build_transcript(segments):
    lines = []

    for segment in segments:
        text = segment.get("text", "").strip()

        if not text:
            continue

        start = format_display_time(
            segment.get("start", 0)
        )

        end = format_display_time(
            segment.get("end", 0)
        )

        speaker = segment.get("speaker", "UNKNOWN")

        lines.append(
            f"[{start} - {end}] {speaker}: {text}"
        )

    return "\n".join(lines)

def get_speaker_statistics(segments):
    stats = {}

    for segment in segments:
        speaker = segment.get("speaker", "UNKNOWN")
        start = float(segment.get("start", 0))
        end = float(segment.get("end", 0))

        duration = max(0, end - start)

        if speaker not in stats:
            stats[speaker] = {
                "segments": 0,
                "duration": 0.0
            }

        stats[speaker]["segments"] += 1
        stats[speaker]["duration"] += duration

    return stats

# =========================================================
# PARSE EDITED TRANSCRIPT
# =========================================================

def parse_edited_transcript(
    transcript,
    original_segments
):

    lines = [
        line.strip()
        for line in transcript.splitlines()
        if line.strip()
    ]

    parsed = []

    timestamp_pattern = re.compile(
        r"^\[(\d{2}:\d{2}(?::\d{2})?)\s*-\s*"
        r"(\d{2}:\d{2}(?::\d{2})?)\]\s*(.*)$"
    )


    def time_to_seconds(value):

        parts = [
            int(x)
            for x in value.split(":")
        ]

        if len(parts) == 2:

            minutes, seconds = parts

            return (
                minutes * 60
                + seconds
            )

        hours, minutes, seconds = parts

        return (
            hours * 3600
            + minutes * 60
            + seconds
        )


    for index, line in enumerate(lines):

        match = timestamp_pattern.match(line)

        if match:

            start = time_to_seconds(
                match.group(1)
            )

            end = time_to_seconds(
                match.group(2)
            )

            text = match.group(3).strip()

            parsed.append({
                "start": start,
                "end": end,
                "text": text
            })

        elif index < len(original_segments):

            original = original_segments[index]

            parsed.append({
                "start": original.get(
                    "start",
                    0
                ),
                "end": original.get(
                    "end",
                    0
                ),
                "text": line
            })

        else:

            last_end = (
                parsed[-1]["end"]
                if parsed
                else 0
            )

            parsed.append({
                "start": last_end,
                "end": last_end + 3,
                "text": line
            })

    return parsed


# =========================================================
# TXT
# =========================================================

def create_txt(text):

    return text.encode(
        "utf-8"
    )


# =========================================================
# DOCX
# =========================================================

def create_docx(text):

    document = Document()

    document.add_heading(
        "Video Transcript",
        level=1
    )

    for line in text.splitlines():

        if line.strip():

            document.add_paragraph(
                line
            )

    fd, temp_path = tempfile.mkstemp(
        suffix=".docx"
    )

    try:

        os.close(fd)

        document.save(
            temp_path
        )

        with open(
            temp_path,
            "rb"
        ) as file:

            return file.read()

    finally:

        if os.path.exists(
            temp_path
        ):

            os.unlink(
                temp_path
            )


# =========================================================
# PDF
# =========================================================

def create_pdf(text):

    fd, temp_path = tempfile.mkstemp(
        suffix=".pdf"
    )

    os.close(fd)

    try:

        document = SimpleDocTemplate(
            temp_path,
            pagesize=A4,
            rightMargin=40,
            leftMargin=40,
            topMargin=40,
            bottomMargin=40
        )

        styles = getSampleStyleSheet()

        title_style = styles["Title"]

        title_style.alignment = TA_CENTER

        body_style = styles["BodyText"]

        body_style.leading = 15

        story = []

        story.append(
            Paragraph(
                "Video Transcript",
                title_style
            )
        )

        story.append(
            Spacer(1, 20)
        )

        for line in text.splitlines():

            if line.strip():

                safe_line = (
                    line
                    .replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                )

                story.append(
                    Paragraph(
                        safe_line,
                        body_style
                    )
                )

                story.append(
                    Spacer(1, 6)
                )

        document.build(story)

        with open(
            temp_path,
            "rb"
        ) as file:

            return file.read()

    finally:

        if os.path.exists(
            temp_path
        ):

            os.unlink(
                temp_path
            )


# =========================================================
# SRT
# =========================================================

def create_srt(segments):

    output = []

    counter = 1

    for segment in segments:

        text = segment.get(
            "text",
            ""
        ).strip()

        if not text:
            continue

        start = format_srt_time(
            segment.get(
                "start",
                0
            )
        )

        end = format_srt_time(
            segment.get(
                "end",
                0
            )
        )

        output.append(
            f"{counter}\n"
            f"{start} --> {end}\n"
            f"{text}\n"
        )

        counter += 1

    return "\n".join(
        output
    ).encode("utf-8")


# =========================================================
# VTT
# =========================================================

def create_vtt(segments):

    output = [
        "WEBVTT",
        ""
    ]

    for segment in segments:

        text = segment.get(
            "text",
            ""
        ).strip()

        if not text:
            continue

        start = format_vtt_time(
            segment.get(
                "start",
                0
            )
        )

        end = format_vtt_time(
            segment.get(
                "end",
                0
            )
        )

        output.append(
            f"{start} --> {end}"
        )

        output.append(text)

        output.append("")

    return "\n".join(
        output
    ).encode("utf-8")


# =========================================================
# JSON
# =========================================================

def create_json(
    segments,
    language,
    model,
    duration
):

    data = {

        "language": language,

        "model": model,

        "duration_seconds": duration,

        "segments": segments
    }

    return json.dumps(
        data,
        ensure_ascii=False,
        indent=4
    ).encode("utf-8")


# =========================================================
# COPY BUTTON
# =========================================================

def copy_button(text):
    import json
    import streamlit.components.v1 as components

    safe_text = json.dumps(text)

    html = f"""
    <button
        id="copyBtn"
        style="
            width:100%;
            padding:12px;
            border:none;
            border-radius:8px;
            background:#ff4b4b;
            color:white;
            font-size:16px;
            font-weight:600;
            cursor:pointer;
        "
    >
        📋 Copy Transcript
    </button>

    <script>
        const textToCopy = {safe_text};
        const button = document.getElementById("copyBtn");

        button.addEventListener("click", async () => {{
            try {{
                if (navigator.clipboard && window.isSecureContext) {{
                    await navigator.clipboard.writeText(textToCopy);
                }} else {{
                    const textarea = document.createElement("textarea");
                    textarea.value = textToCopy;
                    textarea.style.position = "fixed";
                    textarea.style.opacity = "0";
                    document.body.appendChild(textarea);
                    textarea.focus();
                    textarea.select();
                    document.execCommand("copy");
                    textarea.remove();
                }}

                button.innerText = "✅ Copied!";

                setTimeout(() => {{
                    button.innerText = "📋 Copy Transcript";
                }}, 2000);

            }} catch (error) {{
                button.innerText = "❌ Copy failed";

                setTimeout(() => {{
                    button.innerText = "📋 Copy Transcript";
                }}, 2000);
            }}
        }});
    </script>
    """

    components.html(
        html,
        height=60
    )


# =========================================================
# SIDEBAR
# =========================================================

st.sidebar.header(
    "⚙️ Transcription Settings"
)


model_label = st.sidebar.selectbox(
    "Whisper Model",
    list(MODEL_OPTIONS.keys()),
    index=0
)

model_name = MODEL_OPTIONS[
    model_label
]


language_label = st.sidebar.selectbox(
    "Video Language",
    list(LANGUAGE_OPTIONS.keys())
)

language_code = LANGUAGE_OPTIONS[
    language_label
]


translation_enabled = st.sidebar.checkbox(
    "🌐 Translate Transcript",
    value=False
)

TRANSLATION_LANGUAGES = {
    "English": "en",
    "Hindi": "hi",
    "Urdu": "ur",
    "Bengali": "bn",
    "Marathi": "mr",
    "Gujarati": "gu",
    "Tamil": "ta",
    "Telugu": "te",
    "Kannada": "kn",
    "Malayalam": "ml",
    "Punjabi": "pa",
    "Arabic": "ar",
    "Spanish": "es",
    "French": "fr",
    "German": "de",
    "Portuguese": "pt",
    "Russian": "ru",
    "Japanese": "ja",
    "Korean": "ko",
    "Chinese": "zh-CN"
}

if translation_enabled:
    target_language_label = st.sidebar.selectbox(
        "🎯 Translate To",
        list(TRANSLATION_LANGUAGES.keys()),
        index=0
    )

    target_language_code = TRANSLATION_LANGUAGES[target_language_label]
else:
    target_language_label = None
    target_language_code = None

speaker_count_option = st.sidebar.selectbox(
    "👥 Speaker Count",
    ["Auto Detect", "1 Speaker", "2 Speakers", "3 Speakers", "4 Speakers"],
    index=0
)

if speaker_count_option == "Auto Detect":
    num_speakers = None
else:
    num_speakers = int(speaker_count_option.split()[0])

st.sidebar.divider()

st.sidebar.write(
    f"**Processing Device:** {device.upper()}"
)

st.sidebar.write(
    f"**Selected Model:** {model_name}"
)


if model_name == "medium":

    st.sidebar.warning(
        "⚠️ Medium may require significant GPU memory."
    )


# =========================================================
# UPLOAD
# =========================================================

uploaded_file = st.file_uploader(
    "🎥 Upload your video",
    type=[
        "mp4",
        "mov",
        "avi",
        "mkv",
        "webm",
        "mpeg",
        "mpg"
    ]
)


# =========================================================
# GENERATE BUTTON
# =========================================================

if uploaded_file:

    st.success(
        f"Selected: **{uploaded_file.name}**"
    )

    file_size_mb = (
        uploaded_file.size
        / (1024 * 1024)
    )

    st.write(
        f"📦 File size: "
        f"**{file_size_mb:.2f} MB**"
    )


    if st.button(
        "🚀 Generate Transcript",
        type="primary",
        use_container_width=True
    ):

        temp_video = None

        try:

            # -------------------------------------------------
            # SAVE VIDEO
            # -------------------------------------------------

            suffix = Path(
                uploaded_file.name
            ).suffix

            temp_video = tempfile.NamedTemporaryFile(
                delete=False,
                suffix=suffix
            )

            temp_video.write(
                uploaded_file.getbuffer()
            )

            temp_video.close()


            # -------------------------------------------------
            # DURATION
            # -------------------------------------------------

            duration = get_video_duration(
                temp_video.name
            )


            # -------------------------------------------------
            # MODEL
            # -------------------------------------------------

            model = load_whisper_model(
                model_name,
                device
            )


            # -------------------------------------------------
            # OPTIONS
            # -------------------------------------------------

            options = {

                "fp16": False,

                "verbose": False
            }


            if language_code:

                options["language"] = (
                    language_code
                )


            if translation_enabled:
                options["task"] = "transcribe"


            # -------------------------------------------------
            # TRANSCRIBE
            # -------------------------------------------------

            with st.spinner(
                "🎙️ Transcribing video..."
            ):

                result = model.transcribe(
                    temp_video.name,
                    **options,
                    temperature=0,
                    condition_on_previous_text=False,
                    beam_size=1,
                    best_of=1,
                )


            segments = result.get(
                "segments",
                []
            )


            detected_language = result.get(
                "language",
                "unknown"
            )

            # -------------------------------------------------
            # SPEAKER DIARIZATION
            # -------------------------------------------------
            temp_audio = None

            try:
                with st.spinner("👥 Detecting speakers..."):

                    temp_audio = tempfile.NamedTemporaryFile(
                        delete=False,
                        suffix=".wav"
                    )

                    temp_audio.close()

                    # Extract audio from video using FFmpeg
                    subprocess.run(
                        [
                            "ffmpeg",
                            "-y",
                            "-i",
                            temp_video.name,
                            "-vn",
                            "-ac",
                            "1",
                            "-ar",
                            "16000",
                            temp_audio.name
                        ],
                        stdout=subprocess.DEVNULL,
                        stderr=subprocess.DEVNULL,
                        check=True
                    )

                    # Run speaker diarization
                    diarization = perform_speaker_diarization(
                        temp_audio.name,
                        num_speakers=num_speakers
                    )

                    # Attach speaker labels to Whisper segments
                    segments = add_speakers_to_segments(
                        segments,
                        diarization
                    )

                    st.success("👥 Speaker detection completed.")

            except Exception as e:
                st.error("❌ Speaker diarization error:")
                st.exception(e)
                segments = segments

            finally:
                if (
                    temp_audio
                    and os.path.exists(temp_audio.name)
                ):
                    os.unlink(temp_audio.name)
            
            # ============================================================
            # TRANSLATION
            # ============================================================

            translated_segments = []
            translation_error = None

            if not segments:
                st.error(
                    "❌ No spoken content could be detected."
                )
            else:

                # ---------------------------------------------------------
                # TRANSLATE IF ENABLED
                # ---------------------------------------------------------

                if translation_enabled and target_language_code:

                    try:

                        with st.spinner(
                            f"🌐 Translating transcript to {target_language_label}..."
                        ):

                            translator = GoogleTranslator(
                                source="auto",
                                target=target_language_code
                            )

                            def safe_translate(text, retries=3):

                                if not text:
                                    return ""

                                for attempt in range(retries):

                                    try:

                                        translated = translator.translate(text)

                                        if translated is None:
                                            continue

                                        translated = str(translated).strip()

                                        # Reject Google server error messages
                                        error_patterns = [
                                            "Error 500",
                                            "Server Error",
                                            "There was an error",
                                            "Please try again later",
                                            "That's all we know"
                                        ]

                                        if any(
                                            pattern.lower() in translated.lower()
                                            for pattern in error_patterns
                                        ):
                                            continue

                                        if translated:
                                            return translated

                                    except Exception:
                                        continue

                                # If translation completely fails,
                                # keep the original transcript text
                                return text

                            # Translate each segment safely
                            for segment in segments:

                                original_text = (
                                    segment.get("text", "").strip()
                                )

                                translated_text = safe_translate(
                                    original_text
                                )

                                translated_segments.append({
                                    "start": segment.get("start", 0),
                                    "end": segment.get("end", 0),
                                    "text": translated_text,
                                    "speaker": segment.get("speaker", "UNKNOWN")
                                })

                    except Exception as e:

                        translation_error = str(e)

                        translated_segments = []

                # ---------------------------------------------------------
                # CHOOSE WHAT TO DISPLAY
                # ---------------------------------------------------------

                if (
                    translation_enabled
                    and translated_segments
                ):

                    display_segments = (
                        translated_segments
                    )

                    transcript = build_transcript(
                        display_segments
                    )

                else:

                    display_segments = segments

                    transcript = build_transcript(
                        display_segments
                    )

                # ---------------------------------------------------------
                # SHOW TRANSLATION ERROR IF ANY
                # ---------------------------------------------------------

                if translation_error:

                    st.warning(
                        "⚠️ Translation failed. "
                        "Showing original transcript instead."
                    )

                    st.caption(
                        f"Translation error: {translation_error}"
                    )

                # ---------------------------------------------------------
                # SAVE RESULT
                # ---------------------------------------------------------

                st.session_state.generated = True

                st.session_state.transcript = transcript

                st.session_state.edited_transcript = transcript

                # Reset transcript editor with newly generated/translated text
                st.session_state["transcript_editor"] = transcript

                st.session_state.segments = (
                    display_segments
                )

                st.session_state.detected_language = (
                    detected_language
                )

                st.session_state.model_name = (
                    model_name
                )

                st.session_state.duration = (
                    duration
                )

                st.session_state.video_name = (
                    uploaded_file.name
                )

                st.session_state.translation_enabled = (
                    translation_enabled
                )

                st.session_state.target_language = (
                    target_language_label
                )

                st.session_state.translated_segments = (
                    translated_segments
                )

                st.success(
                    "✅ Transcript generated successfully!"
                )


        except Exception as e:

            st.error(
                "❌ Something went wrong."
            )

            st.exception(e)


        finally:

            if (
                temp_video
                and os.path.exists(
                    temp_video.name
                )
            ):

                os.unlink(
                    temp_video.name
                )


# =========================================================
# RESULTS
# IMPORTANT:
# This section is OUTSIDE the Generate button.
# =========================================================

if st.session_state.generated:

    st.divider()

    st.subheader(
        "📊 Transcription Information"
    )


    col1, col2, col3, col4, col5 = st.columns(5)


    with col1:

        st.metric(
            "Language",
            st.session_state.detected_language.upper()
        )


    with col2:

        st.metric(
            "Model",
            st.session_state.model_name
        )


    with col3:

        st.metric(
            "Segments",
            len(st.session_state.segments)
        )


    with col4:

        st.metric(
            "Device",
            device.upper()
        )


    with col5:

        duration = (
            st.session_state.duration
        )

        st.metric(
            "Duration",
            format_display_time(duration)
            if duration
            else "N/A"
        )

    # =====================================================
    # SPEAKER STATISTICS
    # =====================================================

    speaker_stats = get_speaker_statistics(
        st.session_state.segments
    )

    if speaker_stats:
        st.subheader("👥 Speaker Statistics")

        stat_columns = st.columns(len(speaker_stats))

        for column, (speaker, data) in zip(
            stat_columns,
            speaker_stats.items()
        ):
            with column:
                st.metric(
                    speaker,
                    f"{data['duration']:.1f}s"
                )

                st.caption(
                    f"{data['segments']} segment(s)"
                )

    # =====================================================
    # EDITOR
    # =====================================================

    st.subheader(
        "✏️ Editable Transcript"
    )

    edited_transcript = st.text_area(
        "Edit the transcript if required:",
        value=st.session_state.edited_transcript,
        height=450,
        key="transcript_editor"
    )


    # Save edits continuously in session state

    st.session_state.edited_transcript = (
        edited_transcript
    )


    # =====================================================
    # COPY
    # =====================================================

    st.subheader(
        "📋 Copy Transcript"
    )

    copy_button(
        st.session_state.edited_transcript
    )


    # =====================================================
    # SEARCH
    # =====================================================

    st.subheader(
        "🔍 Search Transcript"
    )

    search_term = st.text_input(
        "Search for a word or phrase:",
        placeholder="Example: machine learning",
        key="search_input"
    )


    if search_term.strip():

        lines = (
            st.session_state
            .edited_transcript
            .splitlines()
        )

        matches = [
            line
            for line in lines
            if search_term.lower()
            in line.lower()
        ]


        if matches:

            st.success(
                f"Found {len(matches)} matching line(s)."
            )

            st.text_area(
                "Search Results",
                "\n".join(matches),
                height=250,
                key="search_results"
            )

        else:

            st.warning(
                "No matching text found."
            )


    # =====================================================
    # PREPARE EDITED SEGMENTS
    # =====================================================

    edited_segments = parse_edited_transcript(
        st.session_state.edited_transcript,
        st.session_state.segments
    )


    # =====================================================
    # DOWNLOAD DATA
    # =====================================================

    txt_data = create_txt(
        st.session_state.edited_transcript
    )

    docx_data = create_docx(
        st.session_state.edited_transcript
    )

    pdf_data = create_pdf(
        st.session_state.edited_transcript
    )

    srt_data = create_srt(
        edited_segments
    )

    vtt_data = create_vtt(
        edited_segments
    )

    json_data = create_json(
        edited_segments,
        st.session_state.detected_language,
        st.session_state.model_name,
        st.session_state.duration
    )


    # =====================================================
    # DOWNLOADS
    # =====================================================

    st.subheader(
        "⬇️ Download Transcript"
    )


    col1, col2, col3 = st.columns(3)


    with col1:

        st.download_button(
            "📄 TXT",
            data=txt_data,
            file_name="transcript.txt",
            mime="text/plain",
            use_container_width=True
        )


    with col2:

        st.download_button(
            "📝 DOCX",
            data=docx_data,
            file_name="transcript.docx",
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            use_container_width=True
        )


    with col3:

        st.download_button(
            "📕 PDF",
            data=pdf_data,
            file_name="transcript.pdf",
            mime="application/pdf",
            use_container_width=True
        )


    col4, col5, col6 = st.columns(3)


    with col4:

        st.download_button(
            "📺 SRT",
            data=srt_data,
            file_name="subtitles.srt",
            mime="application/x-subrip",
            use_container_width=True
        )


    with col5:

        st.download_button(
            "🎬 VTT",
            data=vtt_data,
            file_name="subtitles.vtt",
            mime="text/vtt",
            use_container_width=True
        )


    with col6:

        st.download_button(
            "🔢 JSON",
            data=json_data,
            file_name="transcript.json",
            mime="application/json",
            use_container_width=True
        )


    # =====================================================
    # RAW WHISPER DATA
    # =====================================================

    with st.expander(
        "🔎 View Whisper segments"
    ):

        st.json(
            st.session_state.segments
        )